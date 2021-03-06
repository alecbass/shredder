from docx import Document
import os
import sys

def write_sentences_to_docx(fileName: str, sentences: list):
    document = Document()
    for sentence in sentences:
        document.add_paragraph(sentence)
    document.save(fileName[:-5] + "_output.docx")


############### 08/07/2019 Redo

def start(document: Document, wordsToCapture: tuple):
    sentences: list = []
    # NOTE(alec): Sentences
    for para in document.paragraphs:
        for word in wordsToCapture:
            if word in para.text.lower():
                sentences.append(para.text)

    #NOTE(alec): Tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for word in wordsToCapture:
                    if word in cell.text.lower():
                        sentences.append(cell.text)
    return sentences

def paragraphs_to_sentences(paragraphs: list, wordsToCapture: tuple):
    sentences = []
    for para in paragraphs:
        splitPara = para.split(".")
        for word in wordsToCapture:
            for sentence in splitPara:
                if word in sentence:
                    sentences.append(sentence)
    return sentences

def begin_shred(filename: str, wordsToCapture: tuple):
    document = Document(filename)
    validSentences: list = []
    paragraphs = start(document, wordsToCapture)
    sentences = paragraphs_to_sentences(paragraphs, wordsToCapture)
    # tada
    write_sentences_to_docx(filename, sentences)